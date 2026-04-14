from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "outputs" / "doc"
OUTPUT_PATH = OUTPUT_DIR / "paper_outline_smart_home_rule_generation.docx"


TITLE = "中文智能家居自动化规则生成系统的设计与实验研究"

ABSTRACT = (
    "随着智能家居设备规模不断扩大，用户通过自然语言描述自动化规则的需求日益明显，但直接将中文描述映射为可执行配置仍面临结构不稳定、"
    "设备实体易错配、服务调用不一致以及规则缺乏可解释性等问题。针对上述问题，本文设计并实现了一套中文智能家居自动化规则生成系统，"
    "提出以中间领域专用语言（DSL）为核心的规则生成框架，将中文规则描述依次转换为结构化 DSL、经规则校验与自动修复后，再转换为 Home Assistant 风格 YAML，"
    "同时输出中文解释与结构化可视化结果。系统在实现上包含设备知识库、DSL 模型、解析模块、校验模块、修复模块、YAML 转换模块、解释模块和轻量可视化模块，"
    "形成了从输入文本到可执行规则的完整处理链路。为验证方法有效性，本文构建了包含 5000 条设备知识库实体、5000 条中文测试规则以及 1800 条 gold DSL 标注样本的实验数据集，"
    "并从工程稳定性和结构预测准确性两个层面开展实验。结果表明，相较于直接从中文生成 YAML 的基线方案，本文提出的 DSL 中间层方案在 Rule Completeness Rate、"
    "Validation Pass Rate 和 End-to-End Executable Rate 上分别达到 100.00%、100.00% 和 100.00%，显著优于基线方案的 88.20%、70.92% 和 70.92%。"
    "在规则自校验与自动修复实验中，系统将修复前通过率由 30.06% 提升至 99.17%，自动修复增益达到 69.10%，说明“校验 + 局部修复”的组合能够显著增强规则生成的工程稳健性。"
    "在结构化预测实验中，基于 1800 条 gold DSL 的评测结果显示，系统的 Field Accuracy 为 0.5591，Trigger Type F1 为 0.9569，Entity Selection Micro-F1 为 0.8060。"
    "综合来看，本文方法能够有效提升中文智能家居规则生成的结构化程度、规则合法性与最终可执行性，为后续面向真实大模型和真实家庭环境的规则生成研究提供了可复现的实验基础。"
)

KEYWORDS = "智能家居；自动化规则生成；领域专用语言；规则校验；自动修复；可解释性"


SECTIONS = [
    (
        "1 引言",
        [
            "研究背景：说明智能家居设备增多后，用户以自然语言配置自动化规则的需求正在增长。",
            "问题提出：指出直接从中文生成 YAML 容易出现字段缺失、设备错配、服务不兼容和逻辑冲突。",
            "研究目标：提出构建“中文描述 -> DSL -> 校验/修复 -> YAML -> 解释/可视化”的完整系统。",
            "研究意义：强调该工作在规则工程化、可执行性提升和人机交互友好性方面的价值。",
            "本文贡献：概述系统实现、数据集构建、实验设计与结果结论。",
        ],
    ),
    (
        "2 相关研究与技术背景",
        [
            "自然语言到结构化规则：综述从中文或自然语言生成自动化规则、脚本或配置文件的研究思路。",
            "智能家居规则系统：介绍 Home Assistant 类平台中自动化规则的典型结构。",
            "中间表示方法：讨论 DSL 在复杂任务中的结构约束和错误隔离作用。",
            "规则校验与自动修复：总结知识库约束、格式修复、拼写纠错和冲突消解的常见方法。",
            "可解释性与可视化：说明规则解释文本和结构图对用户理解规则的帮助。",
        ],
    ),
    (
        "3 问题定义与总体方法",
        [
            "任务定义：给出输入、输出以及系统目标的形式化描述。",
            "输入输出说明：输入为中文规则描述，输出为 DSL、YAML、解释文本和 Mermaid 图。",
            "总体框架：描述设备知识库、解析器、校验器、修复器、转换器、解释器和可视化器之间的关系。",
            "核心思想：强调 DSL 作为中间层，用于承接语义解析与规则执行之间的结构约束。",
        ],
    ),
    (
        "4 系统设计与实现",
        [
            "设备知识库设计：说明 `devices.json` 的 schema、实体命名规范和服务约束映射。",
            "DSL 设计：介绍 RuleDSL 的字段，包括 `rule_name`、`trigger`、`conditions`、`actions` 和 `mode`。",
            "规则解析模块：说明如何通过可替换 LLM 接口或 mock 模块生成 DSL 草案。",
            "规则校验模块：说明实体存在性、服务兼容性、动作冲突、时间合法性等校验逻辑。",
            "自动修复模块：说明 mode 默认补全、entity 近似匹配、非法 weekday/time 修复等策略。",
            "YAML 转换模块：说明 DSL 向 Home Assistant 风格 YAML 的映射关系。",
            "解释与可视化模块：说明中文解释文本、终端树结构和 Mermaid 图的生成方式。",
        ],
    ),
    (
        "5 数据集构建",
        [
            "设备知识库：介绍 5000 条设备实体的来源、房间覆盖和类型分布。",
            "测试输入集：说明 5000 条中文规则输入及其 normal、ambiguous、error_conflict 三类划分。",
            "gold DSL 标注集：说明 1800 条结构化标注样本的构建原则与人工风格改写策略。",
            "数据质量控制：描述如何过滤明显错误输入、控制房间-设备搭配，并降低与 pipeline 的自一致性。",
            "数据局限：说明当前 gold 仍为高质量启发式标注，未来需进一步人工复核。",
        ],
    ),
    (
        "6 实验设计",
        [
            "实验目标 1：验证引入 DSL 中间层后，相比直接 YAML 生成是否更稳定、更结构化。",
            "实验目标 2：验证规则校验与自动修复是否能提高最终规则通过率。",
            "实验目标 3：验证系统在 gold DSL 上的结构化预测能力。",
            "基线方法：说明直接从中文生成 YAML 的 baseline 设定。",
            "评测指标：列出 Rule Completeness Rate、Validation Pass Rate、End-to-End Executable Rate、Repair Gain、Exact Match、Field Accuracy、Trigger Type F1、Entity Selection Micro-F1。",
            "实验环境：说明 Python 版本、依赖库、本地运行环境以及 mock LLM 配置。",
        ],
    ),
    (
        "7 实验结果与分析",
        [
            "DSL 与直接 YAML 对比：写明 DSL 中间层在三个核心工程指标上优于 baseline。",
            "修复模块结果：分析修复前后通过率变化，并结合错误类型说明修复能力边界。",
            "结构化准确率结果：分析 Exact Match、Field Accuracy、Trigger Type F1 和 Entity Selection Micro-F1。",
            "误差分析：从设备实体 grounding、长尾设备、规则名生成和条件识别等角度分析误差来源。",
            "结果讨论：说明为什么 DSL 中间层可以提升稳定性，但也可能引入新的结构约束偏差。",
        ],
    ),
    (
        "8 论文结果表述建议",
        [
            "建议表 1：DSL 中间层与直接 YAML 基线在三个工程指标上的对比表。",
            "建议表 2：修复前后通过率与不同错误类型修复成功率。",
            "建议表 3：结构化机器学习指标结果表。",
            "建议图 1：系统总体流程图。",
            "建议图 2：单条规则的 Mermaid 结构图示例。",
            "建议图 3：错误类型分布或修复增益可视化。",
        ],
    ),
    (
        "9 结论与未来工作",
        [
            "结论总结：概括本文在结构稳定性、规则合法性和可执行性提升方面的实验结论。",
            "方法价值：强调 DSL 中间层与校验修复组合的工程价值。",
            "当前不足：说明 gold 标注质量、长尾设备覆盖和真实大模型接入仍有限。",
            "未来工作：提出引入人工审校标注、真实 LLM、Hit@K 或 MRR 检索指标以及真实 Home Assistant 联调。",
        ],
    ),
]


def configure_styles(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(12)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(TITLE)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)
    run.bold = True

    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("论文写作提纲与内容框架（实验版）")


def add_abstract(document: Document) -> None:
    document.add_heading("摘要", level=1)
    paragraph = document.add_paragraph(ABSTRACT)
    paragraph.paragraph_format.first_line_indent = Pt(24)

    keyword_para = document.add_paragraph()
    keyword_para.add_run("关键词：").bold = True
    keyword_para.add_run(KEYWORDS)


def add_sections(document: Document) -> None:
    for title, bullets in SECTIONS:
        document.add_heading(title, level=1)
        intro = document.add_paragraph()
        intro.add_run("内容框架：").bold = True
        intro.add_run("建议本节按照以下逻辑展开撰写。")
        for bullet in bullets:
            document.add_paragraph(bullet, style="List Bullet")


def add_appendix(document: Document) -> None:
    document.add_heading("附录：当前实验核心指标（可直接写入论文结果部分）", level=1)
    items = [
        "DSL vs 直接 YAML：Rule Completeness Rate 100.00% vs 88.20%，Validation Pass Rate 100.00% vs 70.92%，End-to-End Executable Rate 100.00% vs 70.92%。",
        "批量实验：总样本数 5000，修复前校验通过率 70.40%，修复后校验通过率 100.00%，端到端可执行率 100.00%，Repair Gain 29.60%。",
        "修复实验：修复前通过率 30.06%，修复后通过率 99.17%，Repair Gain 69.10%。",
        "结构化指标：Sample Count 1800，Exact Match 0.00%，Field Accuracy 0.5591，Trigger Type F1 0.9569，Entity Selection Micro-F1 0.8060。",
    ]
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_document() -> Document:
    document = Document()
    configure_styles(document)

    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    add_title(document)
    document.add_paragraph()
    add_abstract(document)
    document.add_section(WD_SECTION.NEW_PAGE)
    add_sections(document)
    add_appendix(document)
    return document


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
